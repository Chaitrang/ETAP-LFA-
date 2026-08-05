"""
word_writer.py
Opens the bundled Bus_ID.docx template, populates its table with extracted
bus rows + computed Remarks, and saves the result. The template's existing
formatting (fonts, borders, spacing, column widths, alignment) is preserved
exactly - we only ever insert text into cells, and clone the template's own
row XML (deep copy) when more data rows are needed than the template ships
with, so cloned rows are pixel-identical to the original ones.
"""
from __future__ import annotations

import copy
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn

# Column order expected in the template's single table, left to right.
COLUMNS = ["bus_id", "nominal", "voltage_pct", "kw_loading", "amp_loading", "remarks"]


def _set_cell_text(cell, text: str) -> None:
    """Write text into a cell's first paragraph without touching tcPr
    (borders/shading/width/vAlign) or paragraph-level formatting."""
    paragraph = cell.paragraphs[0]
    # Clear any existing runs (template cells are empty, but be defensive)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(str(text))


def _clone_row(table, template_tr):
    """Deep-copy the given <w:tr> XML element and append it to the table,
    so the new row has identical borders/shading/column widths/vAlign to
    the template's own rows. `template_tr` must be a pristine (unmodified)
    row element - callers should capture it once before the fill loop,
    since rows are cloned repeatedly."""
    new_tr = copy.deepcopy(template_tr)
    table._tbl.append(new_tr)
    from docx.table import _Row
    return _Row(new_tr, table)


def populate_template(
    template_path: str,
    output_path: str,
    bus_rows: Iterable[dict],
) -> str:
    """
    bus_rows: iterable of dicts with keys:
        bus_id, nominal, voltage_pct, kw_loading, amp_loading, remarks
    (see main.py for how these are built from pdf_parser.BusRecord +
    voltage_checker.evaluate_remark)

    Returns the output_path on success. Raises on I/O errors (caller
    surfaces the standard "Unable to read the ETAP report."-style message
    per the spec only for PDF read failures - template errors are a
    packaging/build issue, not a user-facing PDF error).
    """
    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("Bundled Word template has no table to populate.")
    table = doc.tables[0]

    bus_rows = list(bus_rows)
    header_row_count = 1
    available_data_rows = len(table.rows) - header_row_count

    # Capture a pristine copy of the template's last row BEFORE writing any
    # text, since it's used as the formatting source for any cloned rows
    # further down (and by then the original rows will already contain text).
    pristine_tr = copy.deepcopy(table.rows[-1]._tr)

    # Fill existing empty data rows first
    for i in range(min(available_data_rows, len(bus_rows))):
        row = table.rows[header_row_count + i]
        data = bus_rows[i]
        for col_idx, key in enumerate(COLUMNS):
            _set_cell_text(row.cells[col_idx], data[key])

    # Clone additional rows (formatted identically to the template's rows)
    # for any buses beyond what the template pre-built.
    if len(bus_rows) > available_data_rows:
        for data in bus_rows[available_data_rows:]:
            row = _clone_row(table, pristine_tr)
            for col_idx, key in enumerate(COLUMNS):
                _set_cell_text(row.cells[col_idx], data[key])

    # If the template shipped with MORE rows than we have buses, blank the
    # unused trailing rows rather than leaving stray placeholder rows with
    # leftover formatting-only content (they are already empty by default,
    # so nothing to do - template ships with truly empty cells).

    doc.save(output_path)
    return output_path

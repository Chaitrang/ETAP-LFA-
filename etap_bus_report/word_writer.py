"""
word_writer.py
--------------
Population of the bundled Word template with the extracted bus data.

The template is never re-created: it is opened, its existing table is filled
cell by cell and the result is saved under a new name, so all fonts, borders,
shading, column widths, header repeat and spacing are preserved exactly.

Blank rows of the template are re-used first.  If the report has more buses
than the template has rows, extra rows are produced by deep-copying the first
blank data row (which carries all of its formatting with it); if it has fewer,
the unused rows are removed.
"""

from __future__ import annotations

import copy
import os
from typing import Iterable, Optional, Sequence

from docx import Document
from docx.table import Table, _Row

from utils import BusRecord, template_path

#: Column order of the template.
COLUMNS = ("Bus ID", "Nominal (kV, A)", "Voltage %", "kW Loading", "Amp Loading", "Remarks")


class TemplateError(Exception):
    """Raised when the bundled template is missing or has an unexpected layout."""


# --------------------------------------------------------------------------- #
# Low level table helpers
# --------------------------------------------------------------------------- #


def _bus_table(document: Document) -> Table:
    """Return the bus table of the template (the first table with 6 columns)."""
    for table in document.tables:
        if len(table.columns) >= len(COLUMNS):
            return table
    raise TemplateError("The bundled Word template does not contain the bus table.")


def _set_cell_text(cell, text: str) -> None:
    """
    Write *text* into *cell*, keeping the cell's own paragraph formatting.

    Only the text is touched: the paragraph properties (alignment, spacing,
    style) and the cell properties (borders, shading, vertical alignment,
    width) defined in the template are left untouched.
    """
    # Keep the first paragraph, drop any others the template may hold.
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)

    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def _append_row(table: Table, prototype: _Row) -> _Row:
    """Append a deep copy of *prototype* (formatting included) to *table*."""
    new_element = copy.deepcopy(prototype._tr)
    prototype._tr.getparent().append(new_element)
    return table.rows[-1]


def _remove_row(table: Table, row: _Row) -> None:
    row._tr.getparent().remove(row._tr)


def _resize_table(table: Table, data_rows: int, header_rows: int = 1) -> None:
    """Make the table hold exactly *data_rows* data rows."""
    current = len(table.rows) - header_rows
    if current <= 0:
        raise TemplateError("The bundled Word template has no blank data row to copy.")

    prototype = table.rows[header_rows]
    while current < data_rows:
        _append_row(table, prototype)
        current += 1
    while current > data_rows:
        _remove_row(table, table.rows[-1])
        current -= 1


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #


def write_bus_report(
    buses: Sequence[BusRecord],
    output_path: str,
    template: Optional[str] = None,
) -> str:
    """
    Fill the template with *buses* and save the document to *output_path*.

    Parameters
    ----------
    buses:
        Bus records in the order they must appear (i.e. the order of the PDF).
    output_path:
        Destination ``.docx`` path.  Parent folders are created if needed.
    template:
        Optional override of the bundled template (used by tests and by the
        future "user-configurable template" feature).

    Returns
    -------
    str
        The path the document was written to.
    """
    template = template or template_path()
    if not os.path.exists(template):
        raise TemplateError(f"Word template not found: {template}")

    document = Document(template)
    table = _bus_table(document)
    _resize_table(table, len(buses))

    for row, bus in zip(table.rows[1:], buses):
        for cell, value in zip(row.cells, bus.as_row()):
            _set_cell_text(cell, value)

    folder = os.path.dirname(os.path.abspath(output_path))
    if folder:
        os.makedirs(folder, exist_ok=True)
    document.save(output_path)
    return output_path


def export_rows(buses: Iterable[BusRecord]) -> list[list[str]]:
    """
    Plain list-of-lists view of the report (header + data).

    Not used by the Word writer itself; provided so that the planned Excel /
    CSV export can share exactly the same formatting rules.
    """
    return [list(COLUMNS)] + [bus.as_row() for bus in buses]

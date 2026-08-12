"""
dynamic_word_writer.py
----------------------
Builds a Word report from an arbitrary header/row table - no template.

The Load Flow report is no longer tied to a fixed template: its columns come
from the user's selection, so the document is generated here.  It receives the
very same table the preview and the Excel workbook use.

The Short Circuit report still uses its own template through
``shortcircuit_writer``; that path is untouched.
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Optional, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT_NAME = "Arial"
HEADER_FILL = "524689"      # matches the house style of the earlier templates
GRID_COLOUR = "7F7F7F"
BODY_SIZE = Pt(9)
HEADER_SIZE = Pt(9)

#: Landscape once the table gets this wide.
LANDSCAPE_FROM_COLUMNS = 6


def _shade(cell, hex_colour: str) -> None:
    shading = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shading)


def _borders(table) -> None:
    """Single thin border on every edge of every cell."""
    properties = table._tbl.tblPr
    borders = properties.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.makeelement(qn(f"w:{edge}"), {})
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), GRID_COLOUR)
        borders.append(element)
    properties.append(borders)


def _repeat_header(row) -> None:
    """Mark a row as a header row so Word repeats it on every page."""
    properties = row._tr.get_or_add_trPr()
    header = properties.makeelement(qn("w:tblHeader"), {})
    header.set(qn("w:val"), "true")
    properties.append(header)


def _write(cell, text: str, bold: bool = False, colour: Optional[str] = None,
           align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = HEADER_SIZE if bold else BODY_SIZE
    run.font.bold = bold
    if colour:
        run.font.color.rgb = RGBColor.from_string(colour)


def build_document(
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    title: str = "LOAD FLOW ANALYSIS REPORT",
    metadata: Optional[dict] = None,
    notes: Optional[Sequence[str]] = None,
) -> Document:
    """
    Build the report document.

    Parameters
    ----------
    headers, rows:
        The final table - exactly what the preview and the workbook show.
    title:
        Report heading.
    metadata:
        ``{label: value}`` lines printed under the title (source file, date,
        limits applied ...).  Only what the caller actually knows.
    notes:
        Footnotes, e.g. columns ETAP could not fill.
    """
    document = Document()

    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    section = document.sections[0]
    if len(headers) >= LANDSCAPE_FROM_COLUMNS:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(title)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True

    for label, value in (metadata or {}).items():
        if value in (None, ""):
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(f"{label}: ")
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = FONT_NAME
        run = paragraph.add_run(str(value))
        run.font.size = Pt(9)
        run.font.name = FONT_NAME

    document.add_paragraph()

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _borders(table)

    header_row = table.rows[0]
    _repeat_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        _write(cell, str(header), bold=True, colour="FFFFFF")
        _shade(cell, HEADER_FILL)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if index >= len(cells):
                break
            align = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _write(cells[index], "" if value is None else str(value), align=align)

    _fit_columns(table, section, headers, rows)

    for note in notes or []:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(note)
        run.font.size = Pt(8)
        run.font.name = FONT_NAME
        run.font.italic = True

    return document


def _fit_columns(table, section, headers, rows) -> None:
    """
    Share the usable page width between the columns in proportion to the
    longest text each one holds, so a short column like Remarks does not get
    squeezed into two lines while a numeric column sits half empty.
    """
    usable = section.page_width - section.left_margin - section.right_margin


    widths = []
    for index, header in enumerate(headers):
        longest = max(
            [len(part) for part in str(header).split()] or [1]
        )
        for row in rows:
            if index < len(row):
                longest = max(longest, len(str(row[index])))
        # Bus IDs are the one column allowed to wrap; cap its demand so a
        # single long ID cannot starve the rest. The constant covers the cell
        # margins, which cost roughly two characters on each side.
        widths.append(min(longest, 22 if index == 0 else 14) + 5)

    total = float(sum(widths)) or 1.0
    resolved = [int(usable * weight / total) for weight in widths]

    # A fixed layout is what makes explicit widths stick; with the default
    # autofit, Word and LibreOffice re-flow the columns to fit their content.
    table.autofit = False
    table.allow_autofit = False

    # The grid is authoritative under a fixed layout, so it has to agree with
    # the cell widths - setting the cells alone has no effect.
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for column, width in zip(grid.findall(qn("w:gridCol")), resolved):
            column.set(qn("w:w"), str(int(width / 635)))       # EMU -> twips

    for index, width in enumerate(resolved):
        for row in table.rows:
            if index < len(row.cells):
                row.cells[index].width = width


def build_document_bytes(headers, rows, title="LOAD FLOW ANALYSIS REPORT",
                         metadata=None, notes=None) -> bytes:
    """The report as raw ``.docx`` bytes."""
    buffer = io.BytesIO()
    build_document(headers, rows, title, metadata, notes).save(buffer)
    return buffer.getvalue()


def default_metadata(source_file: Optional[str] = None, **extra) -> dict:
    """Metadata block from what the application actually knows."""
    metadata = {}
    if source_file:
        metadata["Source report"] = os.path.basename(source_file)
    metadata["Generated"] = datetime.now().strftime("%d %b %Y  %H:%M")
    metadata.update({k: v for k, v in extra.items() if v not in (None, "")})
    return metadata

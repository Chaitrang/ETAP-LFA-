"""
tests.py
--------
Regression tests.  Run with ``pytest tests.py`` (or ``python tests.py`` for a
quick pass/fail summary without pytest).

The PDF-based tests are skipped automatically if no sample report is present in
``samples/``.
"""

from __future__ import annotations

import os
import tempfile

import pdf_parser
import voltage_checker
import word_writer
from utils import BusRecord, fmt_nominal_voltage

_HERE = os.path.dirname(os.path.abspath(__file__))
SAMPLE = os.path.join(_HERE, "samples", "sample_lfa_report.pdf")
SC_SAMPLE = os.path.join(_HERE, "samples", "sample_sc_report.pdf")
HAS_SAMPLE = os.path.exists(SAMPLE)
HAS_SC_SAMPLE = os.path.exists(SC_SAMPLE)


# --------------------------------------------------------------------------- #
# Voltage limits
# --------------------------------------------------------------------------- #


def test_limits_parsing():
    for text, expected in [
        ("95-106", (95.0, 106.0)),
        ("95 - 105 %", (95.0, 105.0)),
        ("90-110", (90.0, 110.0)),
        ("95%-106%", (95.0, 106.0)),
        ("95 to 106", (95.0, 106.0)),
        ("94.5–105.5", (94.5, 105.5)),
    ]:
        limits = voltage_checker.parse_limits(text)
        assert (limits.lower, limits.upper) == expected, text


def test_limits_rejects_bad_input():
    for text in ["95", "95%", "abc", "", "106-95", "95-", "-106"]:
        try:
            voltage_checker.parse_limits(text)
        except voltage_checker.LimitError as exc:
            assert str(exc) == voltage_checker.LIMIT_FORMAT_MESSAGE
        else:  # pragma: no cover
            raise AssertionError(f"{text!r} should have been rejected")


def test_remarks():
    limits = voltage_checker.parse_limits("95-106")
    assert voltage_checker.evaluate(95.0, limits) == "ACCEPTABLE"      # inclusive
    assert voltage_checker.evaluate(106.0, limits) == "ACCEPTABLE"     # inclusive
    assert voltage_checker.evaluate(94.9, limits) == "NOT ACCEPTABLE"
    assert voltage_checker.evaluate(107.8, limits) == "NOT ACCEPTABLE"
    assert voltage_checker.evaluate(None, limits) == ""


def test_nominal_formatting():
    assert fmt_nominal_voltage(132.0) == "132 kV"
    assert fmt_nominal_voltage(11.0) == "11 kV"
    assert fmt_nominal_voltage(3.3) == "3.3 kV"
    assert fmt_nominal_voltage(3.45) == "3.45 kV"
    assert fmt_nominal_voltage(0.415) == "415 V"
    assert fmt_nominal_voltage(0.38) == "380 V"


# --------------------------------------------------------------------------- #
# Word writer (no PDF needed)
# --------------------------------------------------------------------------- #


def test_word_writer_row_count():
    from docx import Document

    buses = [
        BusRecord(f"BUS-{i}", 11.0, 100.0 + i / 10, 1234.5, 678.9, "ACCEPTABLE")
        for i in range(120)  # more than the blank rows in the template
    ]
    with tempfile.TemporaryDirectory() as folder:
        path = word_writer.write_bus_report(buses, os.path.join(folder, "out.docx"))
        table = Document(path).tables[0]
        assert len(table.rows) == len(buses) + 1                     # + header
        assert table.rows[1].cells[0].text == "BUS-0"
        assert table.rows[-1].cells[0].text == "BUS-119"
        assert table.rows[1].cells[1].text == "11 kV"

    # fewer buses than blank template rows -> extra rows removed
    with tempfile.TemporaryDirectory() as folder:
        path = word_writer.write_bus_report(buses[:3], os.path.join(folder, "small.docx"))
        assert len(Document(path).tables[0].rows) == 4


# --------------------------------------------------------------------------- #
# Parser (needs the sample report)
# --------------------------------------------------------------------------- #


def test_parser_sample():
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    buses = pdf_parser.parse_bus_data(SAMPLE, include_pseudo_buses=True)

    # every bus of the report, in report order, no duplicates
    assert len(buses) == 283, len(buses)
    assert len({b.bus_id for b in buses}) == len(buses)
    assert buses[0].bus_id == "132kV GRID"

    by_id = {b.bus_id: b for b in buses}

    # a plain row
    grid = by_id["132kV GRID"]
    assert grid.nominal_text == "132 kV"
    assert grid.voltage_text == "100.0"
    assert grid.amp_loading == 272.8

    # a switchboard: kW must come from the Bus Loading Summary, not from the
    # directly connected load (which is zero for this bus)
    swbd = by_id["S002-SB-1AC001 [BUS A]"]
    assert swbd.nominal_text == "11 kV"
    assert swbd.voltage_text == "101.2"
    assert abs(swbd.kw_loading - 31.398 * 0.910 * 1000) < 1.0
    assert swbd.amp_loading == 1627.6

    # bus IDs wrapped over two lines in the PDF must be reassembled
    assert "S002-SB-PSS01N01 [BUS A]" in by_id
    assert "S003-DB-11N001-03 [BUS B]" in by_id

    # bus IDs that overflow their column must keep every word
    assert by_id["TYPICAL MOTOR FEEDER 1 A"].amp_loading == 1069.5
    assert by_id["POWER SOCKET FEEDER 1 B"].amp_loading == 130.6

    # ETAP internal nodes are dropped by default
    filtered = pdf_parser.parse_bus_data(SAMPLE)
    assert all("~" not in b.bus_id for b in filtered)
    assert len(filtered) < len(buses)


def test_end_to_end():
    """The legacy `limits=` call must classify exactly as it always did."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import cli

    with tempfile.TemporaryDirectory() as folder:
        output = os.path.join(folder, "Bus Report.docx")
        path, result = cli.build_report(SAMPLE, output, "load_flow", limits="95-106")
        assert os.path.exists(path)
        assert result.count == 255
        remarks = {row[-1] for row in result.rows}
        assert remarks <= {"ACCEPTABLE", "NOT ACCEPTABLE", ""}
        # the 107.8 % bus must be flagged, exactly as before
        bus2 = [row for row in result.rows if row[0] == "Bus2"]
        assert bus2 and bus2[-1][-1] == "NOT ACCEPTABLE"
        assert result.flagged == 24


def test_missing_file():
    try:
        pdf_parser.parse_bus_data("does_not_exist.pdf")
    except Exception as exc:
        assert str(exc) == pdf_parser.MSG_UNREADABLE
    else:  # pragma: no cover
        raise AssertionError("missing file should raise")


# --------------------------------------------------------------------------- #
# Short Circuit: template mapping (no PDF needed)
# --------------------------------------------------------------------------- #


def test_template_mapping():
    import shortcircuit_writer
    import template_mapping
    from docx import Document

    mapping = dict(shortcircuit_writer.describe_mapping())
    assert mapping["Switchgear ID"] == "bus_id"
    assert mapping["Bus Rating (kV, A)"] == "nominal_voltage"
    # "Switchgear Rating (kA)" is an equipment rating, not an ID and not a voltage
    assert mapping['Switchgear Rating (kA) Ip (peak)'] == "rating_peak"
    assert mapping['3 Phase ETAP Results (kA) I" k'] == "ik_initial"
    assert mapping["3 Phase ETAP Results (kA) Ip"] == "ip_peak"
    assert mapping["3 Phase ETAP Results (kA) Ik"] == "ik_steady"
    assert mapping["Remarks"] == "remarks"

    # multi-row header with a merged banner row is detected
    table = Document(shortcircuit_writer.template_path()).tables[0]
    assert template_mapping.header_row_count(table) == 3

    # the Load Flow template maps too (same generic machinery)
    from utils import resource_path
    lf = Document(resource_path(os.path.join("assets", "LoadFlow_Template.docx"))).tables[0]
    rows, keys, _ = template_mapping.map_template(lf)
    assert rows == 1
    assert keys == ["bus_id", "nominal_voltage", "voltage_percent",
                    "kw_loading", "amp_loading", "remarks"]


def test_sc_remarks_rules():
    import sc_rules

    ok = {"rating_peak": 100.0, "ip_peak": 45.0, "ik_initial": 19.0,
          "rating_ib_sym": 40.0, "devices": [{"exceeded": False}], "device_exceeded": False}
    bad = {"rating_peak": 40.0, "ip_peak": 84.7, "ik_initial": 49.2,
           "rating_ib_sym": 36.0, "devices": [{"exceeded": True}], "device_exceeded": True}
    unrated = {"rating_peak": None, "ip_peak": 84.7, "devices": [], "device_exceeded": False}

    assert sc_rules.peak_vs_rating(ok) == "ACCEPTABLE"
    assert sc_rules.peak_vs_rating(bad) == "NOT ACCEPTABLE"
    assert sc_rules.peak_vs_rating(unrated) == ""          # no rating -> no claim
    assert sc_rules.peak_and_breaking(ok) == "ACCEPTABLE"
    assert sc_rules.peak_and_breaking(bad) == "NOT ACCEPTABLE"
    assert sc_rules.etap_flag(bad) == "NOT ACCEPTABLE"
    assert sc_rules.etap_flag(unrated) == ""
    assert sc_rules.always_blank(bad) == ""


# --------------------------------------------------------------------------- #
# Short Circuit: parser (needs the sample report)
# --------------------------------------------------------------------------- #


def test_sc_parser_sample():
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import shortcircuit_parser

    buses = shortcircuit_parser.parse_short_circuit(SC_SAMPLE, include_pseudo_buses=True)

    assert len(buses) == 276, len(buses)
    assert len({b["bus_id"] for b in buses}) == len(buses)   # no duplicates
    by_id = {b["bus_id"]: b for b in buses}

    # plain bus: currents from the bus row, rating from the lowest rated device
    grid = by_id["132kV GRID"]
    assert grid["nominal_kv"] == 132.0
    assert (grid["ik_initial"], grid["ip_peak"], grid["ik_steady"]) == (19.094, 45.690, 18.060)
    assert grid["rating_peak"] == 100.0
    assert grid["governing_device"] == "=E02-Q0"

    # bus modelled as switchgear: the assembly's own rating is used
    swgr = by_id["S002-SB-1AC001 [BUS A]"]
    assert swgr["bus_type"] == "Switchgear"
    assert swgr["rating_peak"] == 125.0
    assert swgr["ip_peak"] == 77.725

    # ETAP flagged this bus's breaker with '*'
    assert by_id["200A MDB Agreko"]["device_exceeded"] is True

    # wrapped IDs reassembled, page-break repeats not duplicated,
    # trailing footnotes not glued onto the last ID
    assert "S002-SB-1AD001 [BUS B]" in by_id
    assert "TYPICAL MOTOR FEEDER 1 B" in by_id
    assert all(len(b["bus_id"]) < 60 for b in buses)

    # internal nodes dropped by default
    filtered = shortcircuit_parser.parse_short_circuit(SC_SAMPLE)
    assert all("~" not in b["bus_id"] for b in filtered)


def test_sc_end_to_end():
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    from docx import Document

    import reports

    result = reports.generate("short_circuit", SC_SAMPLE)
    assert result.count == 256
    assert result.document_bytes[:2] == b"PK"

    table = Document(tempfile_bytes(result.document_bytes)).tables[0]
    assert len(table.rows) == result.count + 3          # 3 header rows
    first = [c.text for c in table.rows[3].cells]
    assert first[0] == "132kV GRID"
    assert first[1] == "132 kV"
    assert first[2] == "100.00"
    assert first[6] == "ACCEPTABLE"

    # switchgear-only view
    equipment = reports.generate("short_circuit", SC_SAMPLE, equipment_only=True)
    assert 0 < equipment.count < result.count
    assert all(row[2] for row in equipment.rows)       # every one has a rating


def test_reports_registry():
    import reports

    assert reports.labels() == ["Load Flow Analysis", "Short Circuit Study"]
    assert "limits" in reports.get("load_flow").inputs
    assert reports.get("short_circuit").inputs == ()
    for key in reports.ORDER:
        assert os.path.exists(reports.get(key).template_path())


# --------------------------------------------------------------------------- #
# Configurable Load Flow interface
# --------------------------------------------------------------------------- #


def _records():
    import reports

    if not hasattr(_records, "cache"):
        _records.cache = reports.load_flow_records(SAMPLE)
    return _records.cache


def test_lf_column_selection():
    """Only selected fields appear, in the documented order."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_config as C
    import reports

    # Test 1 - Nominal kV + Voltage
    table = reports.load_flow_table(
        _records(), C.LoadFlowConfig(bus_info=("nominal_kv",), results=("voltage",))
    )
    assert table.headers == ["Bus ID", "Nominal kV", "Voltage (%)", "Remarks"]

    # Test 2 - a wider selection keeps Bus Info before Results
    table = reports.load_flow_table(
        _records(),
        C.LoadFlowConfig(
            bus_info=("nominal_kv", "rated_amp"),
            results=("voltage", "kw_loading", "percent_loading"),
        ),
    )
    assert table.headers == [
        "Bus ID", "Nominal kV", "Amp Rating",
        "Voltage (%)", "kW Loading", "% Loading", "Remarks",
    ]

    # Test 3 - every available field
    import loadflow_fields as F

    table = reports.load_flow_table(
        _records(),
        C.LoadFlowConfig(
            bus_info=tuple(F.available_keys(F.BUS_INFO_ORDER)),
            results=tuple(F.available_keys(F.RESULT_ORDER)),
            include_alert_cause=True,
        ),
    )
    assert table.headers[0] == "Bus ID"
    assert table.headers[-2:] == ["Alert", "Remarks"]
    assert len(table.headers) == 1 + 2 + 6 + 2

    # Test 4 - a single result column is still a valid report
    table = reports.load_flow_table(
        _records(), C.LoadFlowConfig(bus_info=(), results=("voltage",))
    )
    assert table.headers == ["Bus ID", "Voltage (%)", "Remarks"]

    # unselected fields never appear as empty columns
    assert "kvar Loading" not in table.headers


#: kW / kvar as ETAP's own results grid shows them, for buses visible in the
#: reference screenshot. These are the ground truth for extraction accuracy.
ETAP_GRID = {
    # bus                     kW        kvar
    "132kV GRID":          (56579.9, 26252.2),
    "200A MDB Agreko":     (   22.06,   10.67),
    "Area Lighting DB (A8)": (  9.06,    6.79),
    "Bus17":               (  493.7,   349.4),
    "Bus18":               (  564.0,   442.1),
    "Bus28":               (  491.0,   335.7),
    "Bus33":               (  560.3,   422.4),
    "Bus36":               (  646.6,   485.0),
    "Bus11":               (    0.0,     0.001),
    "Bus35":               (    0.0,     0.0037),
    "Bus41":               (    0.0,     0.0044),
    "Bus58":               (    0.0,     0.0045),
    "Bus61":               (    0.0,     0.0012),
}


def test_kw_matches_etap_grid():
    """
    kW / kvar must match ETAP's own results grid to well under 1 kW.

    The report prints %PF to one decimal, so deriving kW as MVA x %PF is out by
    ~9 kW on a 56 MW bus. The parser therefore prefers the Load Flow Report's
    own MW figures (directly connected load + outgoing branch flows), which are
    printed to 3 decimals of MW at any bus size.
    """
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import pdf_parser

    buses = {r.bus_id: r for r in pdf_parser.parse_bus_data(SAMPLE)}

    for bus_id, (kw, kvar) in ETAP_GRID.items():
        record = buses[bus_id]
        assert record.kw_loading is not None, bus_id
        assert abs(record.kw_loading - kw) <= 0.5, (bus_id, record.kw_loading, kw)
        assert abs((record.meta.get("kvar_loading") or 0.0) - kvar) <= 7.0, bus_id

    # every bus gets a value - the old directly-connected-load fallback was
    # unreachable because the Load MW column was never mapped
    assert all(r.kw_loading is not None for r in buses.values())

    # the Generation / Load / Load Flow MW columns must stay distinct
    grid = buses["132kV GRID"]
    assert grid.meta["load_mw"] == 0.0            # the grid supplies, it does not consume
    assert grid.meta["flows_mw"]                  # its export is a branch flow


def test_lf_units_and_voltage_display():
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_config as C
    import reports

    base = reports.load_flow_table(
        _records(), C.LoadFlowConfig(bus_info=("nominal_kv",), results=("voltage", "kw_loading"))
    )
    mva = reports.load_flow_table(
        _records(),
        C.LoadFlowConfig(
            bus_info=("nominal_kv",), results=("voltage", "kw_loading"),
            power_unit="MVA", voltage_display="Actual Value",
        ),
    )
    assert base.headers[2:4] == ["Voltage (%)", "kW Loading"]
    assert mva.headers[2:4] == ["Voltage (kV)", "MW Loading"]

    # 132 kV bus at 100 % -> 132.00 kV, and kW/1000 = MW
    assert base.rows[0][2] == "100.00" and mva.rows[0][2] == "132.00"
    assert abs(float(base.rows[0][3]) / 1000 - float(mva.rows[0][3])) < 0.05


def test_lf_alert_classification():
    """Thresholds drive the status, and legacy settings reproduce the old result."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_config as C
    import reports

    # legacy band 95-106, marginals off -> the pre-existing 24 flagged buses
    legacy = reports.load_flow_table(_records(), C.from_legacy_limits("95-106"))
    assert legacy.status_counts.get("CRITICAL") == 24
    assert legacy.status_counts.get("MARGINAL") is None
    assert {row[-1] for row in legacy.rows} <= {"ACCEPTABLE", "NOT ACCEPTABLE"}

    # tightening overvoltage moves buses into CRITICAL
    tighter = C.from_legacy_limits("95-106")
    tighter.overvoltage.critical = 103.0
    tighter.remark_style = C.STATUS_STYLE
    strict = reports.load_flow_table(_records(), tighter)
    assert strict.status_counts["CRITICAL"] > legacy.status_counts["CRITICAL"]

    # a marginal band creates the middle class without changing the critical one
    with_marginal = C.from_legacy_limits("95-106")
    with_marginal.overvoltage.marginal = 102.0
    with_marginal.overvoltage.marginal_enabled = True
    with_marginal.remark_style = C.STATUS_STYLE
    banded = reports.load_flow_table(_records(), with_marginal)
    assert banded.status_counts["CRITICAL"] == 24
    assert banded.status_counts["MARGINAL"] > 0


def test_lf_config_validation():
    import loadflow_config as C

    def rejects(config):
        try:
            config.validate()
        except C.ConfigError:
            return True
        return False

    assert rejects(C.LoadFlowConfig(bus_info=(), results=()))          # nothing selected
    assert rejects(C.LoadFlowConfig(output_format="PDF"))              # unknown format
    assert rejects(C.LoadFlowConfig(bus_info=("does_not_exist",)))     # unknown field

    # marginal beyond critical is inconsistent, in both directions
    bad_over = C.LoadFlowConfig()
    bad_over.overvoltage = C.AlertLimit(105.0, 110.0)
    assert rejects(bad_over)
    bad_under = C.LoadFlowConfig()
    bad_under.undervoltage = C.AlertLimit(95.0, 90.0)
    assert rejects(bad_under)

    # remarks on with every limit off is rejected rather than silently blank
    no_limits = C.LoadFlowConfig()
    for limit in (no_limits.loading, no_limits.overvoltage, no_limits.undervoltage):
        limit.critical_enabled = limit.marginal_enabled = False
    assert rejects(no_limits)

    assert not rejects(C.LoadFlowConfig())                              # defaults are valid


def test_lf_never_fabricates():
    """Unavailable quantities stay empty and are reported, never invented."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_fields as F

    assert F.get("bus_type").available is False       # not in a Load Flow report

    import loadflow_config as C
    import reports

    table = reports.load_flow_table(
        _records(), C.LoadFlowConfig(bus_info=("nominal_kv", "rated_amp"), results=("voltage",))
    )
    column = table.columns.index("rated_amp")
    blanks = sum(1 for row in table.rows if not row[column])
    assert blanks > 0                                 # ETAP rates only some buses
    assert any("Amp Rating" in warning for warning in table.warnings)


def test_lf_preview_word_excel_identical():
    """Preview, Word and Excel must hold the same table."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    from docx import Document
    from openpyxl import load_workbook

    import loadflow_config as C
    import reports

    config = C.LoadFlowConfig(
        bus_info=("nominal_kv", "rated_amp"),
        results=("voltage", "kw_loading", "kvar_loading", "percent_loading"),
        include_alert_cause=True,
    )
    table = reports.load_flow_table(_records(), config)
    result = reports.load_flow_result(table, config, source_file=SAMPLE)
    result.excel_bytes = reports.build_excel(result, "Load Flow Analysis")

    # preview == result
    assert result.headers == table.headers
    assert result.rows == [list(row) for row in table.rows]

    # Word
    document_table = Document(tempfile_bytes(result.document_bytes)).tables[0]
    assert [c.text for c in document_table.rows[0].cells] == table.headers
    assert len(document_table.rows) == len(table.rows) + 1
    for index in (0, 1, 17, len(table.rows) - 1):
        assert [c.text for c in document_table.rows[index + 1].cells] == table.rows[index]

    # Excel
    sheet = load_workbook(tempfile_bytes(result.excel_bytes)).active
    offset = 3                        # title row + blank row + header row
    assert [c.value for c in sheet[offset]] == table.headers
    assert sheet.max_row == len(table.rows) + offset
    for index in (0, 1, 17, len(table.rows) - 1):
        row = [sheet.cell(row=index + offset + 1, column=c + 1).value
               for c in range(len(table.headers))]
        expected = table.rows[index]
        for got, want in zip(row, expected):
            text = "" if got is None else str(got)
            if want and _NUMERIC(want):
                assert abs(float(text) - float(want)) < 1e-9
            else:
                assert text == want


def test_lf_output_format_selection():
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_config as C
    import reports

    for output_format, has_word, has_excel in (
        (C.WORD, True, False),
        (C.EXCEL, False, True),
        (C.BOTH, True, True),
    ):
        result = reports.generate(
            "load_flow", SAMPLE, config=C.LoadFlowConfig(output_format=output_format)
        )
        assert bool(result.document_bytes) is has_word, output_format
        assert bool(result.excel_bytes) is has_excel, output_format


def test_lf_parses_once():
    """Changing selections must not touch the PDF again."""
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    import loadflow_config as C
    import pdf_parser
    import reports

    records = reports.load_flow_records(SAMPLE)

    calls = {"n": 0}
    original = pdf_parser.parse_bus_data

    def counting(*args, **kwargs):
        calls["n"] += 1
        return original(*args, **kwargs)

    pdf_parser.parse_bus_data = counting
    try:
        for results in (("voltage",), ("voltage", "kw_loading"), ("amp_loading",)):
            reports.load_flow_table(records, C.LoadFlowConfig(results=results))
    finally:
        pdf_parser.parse_bus_data = original

    assert calls["n"] == 0


# --------------------------------------------------------------------------- #
# Excel export
# --------------------------------------------------------------------------- #


def test_excel_matches_word():
    """The workbook must hold exactly the rows written into the document."""
    from docx import Document
    from openpyxl import load_workbook

    import excel_writer
    import reports

    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    result = reports.generate("short_circuit", SC_SAMPLE, equipment_only=True)
    assert result.excel_bytes[:2] == b"PK"
    assert result.excel_filename == "Short Circuit Report.xlsx"

    sheet = load_workbook(tempfile_bytes(result.excel_bytes)).active
    table = Document(tempfile_bytes(result.document_bytes)).tables[0]
    header_rows = 3

    # same shape
    assert sheet.max_row == result.count + 1
    assert sheet.max_column == len(result.headers)
    assert len(table.rows) - header_rows == result.count

    # same values, row by row, against the Word document itself
    for index, doc_row in enumerate(table.rows[header_rows:], start=2):
        for column, cell in enumerate(doc_row.cells, start=1):
            value = sheet.cell(row=index, column=column).value
            text = "" if value is None else str(value)
            expected = cell.text.strip()
            if expected and _NUMERIC(expected):
                assert abs(float(text) - float(expected)) < 1e-9, (index, column)
            else:
                assert text == expected, (index, column, text, expected)

    # styling essentials
    assert sheet.freeze_panes == "A2"
    assert sheet.cell(row=1, column=1).font.bold
    assert sheet.cell(row=2, column=1).border.left.style == "thin"
    assert sheet.cell(row=2, column=3).number_format == "0.00"
    assert sheet.column_dimensions["A"].width >= 9

    # naming helper
    assert excel_writer.excel_path_for("/x/Bus Report.docx") == "/x/Bus Report.xlsx"


def _NUMERIC(text: str) -> bool:
    import re

    return bool(re.match(r"^-?\d+(\.\d+)?$", text))


def test_excel_for_load_flow():
    if not HAS_SAMPLE:
        print("skipped: samples/sample_lfa_report.pdf not present")
        return

    from openpyxl import load_workbook

    import reports

    result = reports.generate("load_flow", SAMPLE, limits="95-106")
    assert result.excel_filename == "LoadFlow_Report.xlsx"
    sheet = load_workbook(tempfile_bytes(result.excel_bytes)).active
    assert [c.value for c in sheet[3]] == result.headers     # title, blank, header
    assert sheet.max_row == result.count + 3
    assert sheet.cell(row=4, column=1).value == "132kV GRID"
    assert sheet.title == "Load Flow Analysis"


# --------------------------------------------------------------------------- #
# Configurable Short Circuit interface
# --------------------------------------------------------------------------- #


def _sc_records():
    import reports

    if not hasattr(_sc_records, "cache"):
        _sc_records.cache = (
            reports.short_circuit_records(SC_SAMPLE),
            reports.short_circuit_study_info(SC_SAMPLE),
        )
    return _sc_records.cache


def test_sc_study_info_detected():
    """Standard, study type and study case are read, never assumed."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    _, info = _sc_records()
    assert info["standard"] == "IEC"
    assert info["study_type"] == "3-phase"
    assert info["study_case"] == "SC_Max"

    import shortcircuit_config as C

    config = C.from_study_info(info)
    assert (config.standard, config.study_type, config.selected_report) == (
        "IEC", "3-phase", "SC_Max",
    )


def test_sc_column_selection():
    """Only selected fields appear, in the documented order."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C

    records, info = _sc_records()

    # the reference selection: ID, Nominal kV, Rated ip, I"k, ip, Ik
    table = reports.short_circuit_table(
        records,
        C.ShortCircuitConfig(info=("nominal_kv", "rating_peak"),
                             results=("ik_initial", "ip_peak", "ik_steady")),
        info,
    )
    assert table.headers == [
        "ID", "Nominal kV", "Rated ip (kA)", 'I"k (kA)', "ip (kA)", "Ik (kA)", "Remarks",
    ]

    # a narrower selection drops the unselected columns entirely
    table = reports.short_circuit_table(
        records, C.ShortCircuitConfig(info=("nominal_kv",), results=("ik_initial",)), info
    )
    assert table.headers == ["ID", "Nominal kV", 'I"k (kA)', "Remarks"]

    # Info always precedes Results, whatever order they were given in
    table = reports.short_circuit_table(
        records,
        C.ShortCircuitConfig(
            info=("bus_type", "nominal_kv"), results=("ik_steady", "ik_initial"),
            include_duty=True,
        ),
        info,
    )
    assert table.headers == [
        "ID", "Nominal kV", "Type", 'I"k (kA)', "Ik (kA)", "Duty (%)", "Remarks",
    ]


def test_sc_units_convert_consistently():
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C

    records, info = _sc_records()
    base_config = dict(info=("nominal_kv", "rating_peak"), results=("ip_peak",))

    ka = reports.short_circuit_table(records, C.ShortCircuitConfig(**base_config), info)
    amps = reports.short_circuit_table(
        records,
        C.ShortCircuitConfig(current_unit="A", voltage_unit="V", **base_config),
        info,
    )
    assert ka.headers == ["ID", "Nominal kV", "Rated ip (kA)", "ip (kA)", "Remarks"]
    assert amps.headers == ["ID", "Nominal V", "Rated ip (A)", "ip (A)", "Remarks"]

    # same bus, exactly 1000x - and the classification must not change
    assert abs(float(amps.rows[0][3]) - float(ka.rows[0][3]) * 1000) < 1e-6
    assert [r[-1] for r in amps.rows] == [r[-1] for r in ka.rows]


def test_sc_alert_thresholds():
    """Thresholds drive the status; 100 % reproduces the original rule."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import sc_rules
    import shortcircuit_config as C
    import shortcircuit_processor as P

    records, info = _sc_records()

    default = reports.short_circuit_table(records, C.ShortCircuitConfig(), info)
    assert default.status_counts.get("CRITICAL") == 2

    # Critical 100 % must agree bus by bus with the pre-existing rule
    config = C.ShortCircuitConfig(marginal_enabled=False)
    strict = reports.short_circuit_table(records, config, info)
    rows = {row[0]: row[-1] for row in strict.rows}
    for record in records:
        if "~" in record["bus_id"] or record["bus_id"] not in rows:
            continue
        legacy = sc_rules.peak_vs_rating(record)
        new = rows[record["bus_id"]]
        expected = {"": "", "ACCEPTABLE": "ACCEPTABLE", "NOT ACCEPTABLE": "CRITICAL"}[legacy]
        assert new == expected, (record["bus_id"], legacy, new)

    # lowering the critical threshold flags more buses
    lower = reports.short_circuit_table(
        records, C.ShortCircuitConfig(critical=60.0, marginal=50.0), info
    )
    assert lower.status_counts["CRITICAL"] > default.status_counts["CRITICAL"]

    # duty % is ip / rated ip
    row = next(r for r in records if r.get("rating_peak") and r.get("ip_peak"))
    assert abs(P.duty_percent(row) - row["ip_peak"] / row["rating_peak"] * 100) < 1e-9


def test_sc_filters():
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C
    import shortcircuit_processor as P

    records, info = _sc_records()

    everything = reports.short_circuit_table(
        records, C.ShortCircuitConfig(skip_nodes=False), info
    )
    without_nodes = reports.short_circuit_table(records, C.ShortCircuitConfig(), info)
    assert without_nodes.count < everything.count
    assert all("~" not in row[0] for row in without_nodes.rows)
    assert without_nodes.filtered_out.get("nodes")

    switchgear = reports.short_circuit_table(
        records, C.ShortCircuitConfig(equipment_types=("Switchgear",), info=("bus_type",)), info
    )
    assert switchgear.count == 30
    assert {row[1] for row in switchgear.rows} == {"Switchgear"}

    alerted = reports.short_circuit_table(
        records, C.ShortCircuitConfig(skip_non_alerted=True), info
    )
    assert alerted.count == 2
    assert all(row[-1] in ("CRITICAL", "MARGINAL") for row in alerted.rows)

    # filters that remove everything must explain themselves, not export blank
    try:
        reports.short_circuit_table(
            records,
            C.ShortCircuitConfig(equipment_types=("Switchgear",), skip_non_alerted=True),
            info,
        )
    except P.EmptyTableError as exc:
        assert "Skip non-alerted" in str(exc)
    else:
        raise AssertionError("an empty table must be refused")


def test_sc_config_validation():
    import shortcircuit_config as C

    def rejects(**kwargs):
        try:
            C.ShortCircuitConfig(**kwargs).validate()
        except C.ConfigError:
            return True
        return False

    assert rejects(info=(), results=())                       # nothing selected
    assert rejects(results=("ib_sym",))                       # not published
    assert rejects(critical=90.0, marginal=95.0)              # marginal above critical
    assert rejects(output_format="PDF")
    assert rejects(current_unit="mA")
    assert rejects(standard="BS")
    assert rejects(equipment_types=("Transformer",))
    assert rejects(critical_enabled=False, marginal_enabled=False)   # remarks with no limits
    assert rejects(critical_enabled=False, marginal_enabled=False,
                   include_remarks=False, skip_non_alerted=True)
    assert not rejects()                                      # defaults are valid


def test_sc_never_fabricates():
    """Unavailable quantities stay empty and are reported, never invented."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C
    import shortcircuit_fields as F

    # per-bus Ib / Idc are not published, so their checkboxes are disabled
    assert F.get("ib_sym").available is False
    assert F.get("idc").available is False

    records, info = _sc_records()
    table = reports.short_circuit_table(
        records, C.ShortCircuitConfig(info=("nominal_kv", "rating_peak")), info
    )
    column = table.columns.index("rating_peak")
    blanks = sum(1 for row in table.rows if not row[column])
    assert blanks > 0                       # only some buses have a rated device
    assert any("Rated ip" in w for w in table.warnings)
    # those rows must have a blank remark, not a guessed one
    assert all(not row[-1] for row in table.rows if not row[column])


def test_sc_preview_word_excel_identical():
    """The reference selection must be identical in preview, Word and Excel."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    from docx import Document
    from openpyxl import load_workbook

    import reports
    import shortcircuit_config as C

    records, info = _sc_records()
    config = C.ShortCircuitConfig(
        info=("nominal_kv", "bus_type", "rating_peak"),
        results=("ik_initial", "ip_peak", "ik_steady"),
        include_duty=True,
    )
    table = reports.short_circuit_table(records, config, info)
    result = reports.short_circuit_result(table, config, source_file=SC_SAMPLE, study_info=info)
    result.excel_bytes = reports.build_excel(result, "Short Circuit")

    assert result.headers == table.headers
    assert result.rows == [list(r) for r in table.rows]
    assert result.default_filename == "ShortCircuit_SC_Max.docx"
    assert result.excel_filename == "ShortCircuit_SC_Max.xlsx"

    document_table = Document(tempfile_bytes(result.document_bytes)).tables[0]
    assert [c.text for c in document_table.rows[0].cells] == table.headers
    assert len(document_table.rows) == len(table.rows) + 1

    sheet = load_workbook(tempfile_bytes(result.excel_bytes)).active
    offset = 3
    assert [c.value for c in sheet[offset]] == table.headers
    assert sheet.max_row == len(table.rows) + offset

    for index in (0, 1, 12, len(table.rows) - 1):
        assert [c.text for c in document_table.rows[index + 1].cells] == table.rows[index]
        row = [sheet.cell(row=index + offset + 1, column=c + 1).value
               for c in range(len(table.headers))]
        for got, want in zip(row, table.rows[index]):
            text = "" if got is None else str(got)
            if want and _NUMERIC(want):
                assert abs(float(text) - float(want)) < 1e-9
            else:
                assert text == want


def test_sc_output_format_selection():
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C

    for output_format, has_word, has_excel in (
        (C.WORD, True, False), (C.EXCEL, False, True), (C.BOTH, True, True),
    ):
        result = reports.generate(
            "short_circuit", SC_SAMPLE,
            config=C.ShortCircuitConfig(output_format=output_format), dynamic=True,
        )
        assert bool(result.document_bytes) is has_word, output_format
        assert bool(result.excel_bytes) is has_excel, output_format


def test_sc_parses_once():
    """Changing selections must not touch the PDF again."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports
    import shortcircuit_config as C
    import shortcircuit_parser

    records, info = _sc_records()
    calls = {"n": 0}
    original = shortcircuit_parser.parse_short_circuit

    def counting(*args, **kwargs):
        calls["n"] += 1
        return original(*args, **kwargs)

    shortcircuit_parser.parse_short_circuit = counting
    try:
        for results in (("ik_initial",), ("ip_peak",), ("ik_initial", "ik_steady")):
            reports.short_circuit_table(records, C.ShortCircuitConfig(results=results), info)
    finally:
        shortcircuit_parser.parse_short_circuit = original

    assert calls["n"] == 0


def test_sc_template_path_unchanged():
    """The original template-filling report must still work exactly as before."""
    if not HAS_SC_SAMPLE:
        print("skipped: samples/sample_sc_report.pdf not present")
        return

    import reports

    result = reports.generate("short_circuit", SC_SAMPLE)
    assert result.count == 256
    assert result.headers[0] == "Switchgear ID"
    assert result.default_filename == "Short Circuit Report.docx"
    assert result.document_bytes[:2] == b"PK"


def tempfile_bytes(data: bytes):
    """Wrap document bytes so python-docx can open them."""
    import io

    return io.BytesIO(data)


if __name__ == "__main__":
    failures = 0
    for name, func in sorted(globals().items()):
        if name.startswith("test_") and callable(func):
            try:
                func()
                print(f"PASS  {name}")
            except AssertionError as exc:
                failures += 1
                print(f"FAIL  {name}: {exc}")
    print("\nAll tests passed." if not failures else f"\n{failures} test(s) failed.")
    raise SystemExit(1 if failures else 0)
